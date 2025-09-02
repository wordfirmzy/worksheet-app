# worksheet_generator.py
import re
import random
import os
from collections import defaultdict
from docx import Document
from fpdf import FPDF

# ---------------------------
# Configuration / regexes
# ---------------------------
CJK_RE = re.compile(r'[\u4e00-\u9fff]')
CHINESE_RE = CJK_RE
CJK_PUNCT_RE = r'！？。；，、：“”‘’（）《》【】'

# Common abbreviations that should not be treated as sentence ends
ABBREVIATIONS = [
    "Mr.", "Mrs.", "Dr.", "St.", "Prof.", "Inc.", "Ltd.", "Jr.", "Sr.", "vs."
]

# ---------------------------
# ASS Subtitle Cleaning
# ---------------------------
def clean_subtitle_line(line: str) -> str:
    if not line.startswith("Dialogue:"):
        return ""
    parts = line.split(",", 9)
    if len(parts) < 10:
        return ""
    text = parts[-1].strip()
    text = re.sub(r"\{.*?\}", "", text)  # remove style codes
    text = text.replace("\\N", " ")
    if re.search(r"[\u4e00-\u9fff]", text):
        return ""
    return text.strip()

def parse_ass_file(file_path: str) -> list[str]:
    sentences = []
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        for line in f:
            cleaned = clean_subtitle_line(line)
            if cleaned:
                sentences.append(cleaned)
    return sentences

# ---------------------------
# SRT Subtitle Cleaning
# ---------------------------
def parse_srt_file(file_path: str) -> list[str]:
    sentences = []
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        raw_text = f.read()

    blocks = raw_text.split("\n\n")
    for block in blocks:
        lines = block.strip().split("\n")
        if len(lines) >= 3:
            text_lines = lines[2:]
            for t in text_lines:
                t = re.sub(r"<[^>]+>", "", t)  # remove HTML tags
                if re.search(r"[\u4e00-\u9fff]", t):
                    continue
                cleaned = clean_sentence(t)
                if cleaned:
                    sentences.append(cleaned)
    return sentences

# ---------------------------
# Generic Subtitle Parsing (ASS / SRT / TXT)
# ---------------------------
def parse_subtitles(file_path, bilingual_mode=False):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".ass":
        return parse_ass_file(file_path)
    if ext == ".srt":
        return parse_srt_file(file_path)

    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        raw_text = f.read()

    lines = raw_text.splitlines()
    sentences = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        line = clean_sentence(line)
        if line:
            sentences.append(line)
    return sentences

# ---------------------------
# Cleaning (non-ASS generic text)
# ---------------------------
def clean_sentence(text):
    text = re.sub(rf'^[{CJK_PUNCT_RE}]+|[{CJK_PUNCT_RE}]+$', '', text)
    for abbr in ABBREVIATIONS:
        text = text.replace(abbr, abbr.replace('.', '§'))
    text = text.replace('§', '.')
    return text.strip()

# ---------------------------
# Dictionary building
# ---------------------------
def build_dictionaries(sentences):
    sentence_dict = {}
    word_index = defaultdict(list)
    freq_dict = defaultdict(int)

    for idx, sent in enumerate(sentences):
        sentence_dict[idx] = sent
        words = re.findall(r'\b\w+\b', sent)
        for w in words:
            w_clean = w.lower()
            word_index[w_clean].append(idx)
            freq_dict[w_clean] += 1

    return sentence_dict, word_index, freq_dict

# ---------------------------
# Word filtering
# ---------------------------
def filter_words(freq_dict, familiarity='once', level='beginner'):
    filtered = []
    for w, freq in freq_dict.items():
        if familiarity == 'once' and freq == 1:
            filtered.append(w)
        elif familiarity == 'twice' and freq == 2:
            filtered.append(w)
        elif familiarity == 'more' and freq > 2:
            filtered.append(w)
    return filtered

# ---------------------------
# PDF / DOCX generation
# ---------------------------
def save_as_pdf(output_path, worksheet, word_bank):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, "Worksheet\n\n")
    for line in worksheet:
        pdf.multi_cell(0, 10, line)
        pdf.ln(2)
    pdf.multi_cell(0, 10, "\nWord Bank:\n" + ", ".join(word_bank))
    pdf.output(output_path)

def save_as_docx(output_path, worksheet, word_bank):
    doc = Document()
    doc.add_heading("Worksheet", 0)
    for line in worksheet:
        doc.add_paragraph(line)
    doc.add_paragraph("\nWord Bank:\n" + ", ".join(word_bank))
    doc.save(output_path)

# ---------------------------
# Worksheet generation (PDF/DOCX)
# ---------------------------
def generate_worksheet(sentence_dict, word_index, word_list, num_words=12, bilingual_mode=False):
    chosen_words = random.sample(word_list, min(num_words, len(word_list)))
    worksheet = []
    word_bank = []
    used_sentences = set()

    for word in chosen_words:
        if word not in word_index:
            continue
        sentence_ids = word_index[word]
        available_sids = [sid for sid in sentence_ids if sid not in used_sentences]
        if not available_sids:
            word_bank.append(word)
            continue
        sid = random.choice(available_sids)
        sentence = sentence_dict[sid]

        pattern = re.compile(rf"(?<!\\w){re.escape(word)}(?!\\w)", re.IGNORECASE)
        blanked = pattern.sub("________________________", sentence)

        worksheet.append(blanked)
        word_bank.append(word)
        used_sentences.add(sid)

    random.shuffle(word_bank)
    return worksheet, word_bank

# ---------------------------
# Interactive Web Output
# ---------------------------
def generate_worksheet_web(sentence_dict, word_index, word_list, num_words=12, bilingual_mode=False):
    chosen_words = random.sample(word_list, min(num_words, len(word_list)))
    worksheet = []
    word_bank = []
    used_sentences = set()

    for word in chosen_words:
        if word not in word_index:
            continue
        sentence_ids = word_index[word]
        available_sids = [sid for sid in sentence_ids if sid not in used_sentences]
        if not available_sids:
            word_bank.append(word)
            continue
        sid = random.choice(available_sids)
        sentence = sentence_dict[sid]

        pattern = re.compile(rf"(?<!\\w){re.escape(word)}(?!\\w)", re.IGNORECASE)
        blanked, count = pattern.subn("_____", sentence, count=1)

        worksheet.append({"sentence": blanked, "answer": word})
        word_bank.append(word)
        used_sentences.add(sid)

    random.shuffle(word_bank)
    return worksheet, word_bank
